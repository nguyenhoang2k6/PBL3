import json
import re
import unicodedata
import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('docx')
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, fill):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def remove_accents(input_str):
    s = re.sub(r'[đĐ]', 'd', input_str)
    nfkd_form = unicodedata.normalize('NFKD', s)
    return u"".join([c for c in nfkd_form if not unicodedata.combining(c)])

def to_snake_case(name):
    if name.lower() in ['id', 'sku', 'category_id', 'customer_id', 'order_id', 'product_id', 'user_id', 'supplier_id']:
        return name.lower()
    clean_name = remove_accents(name)
    clean_name = re.sub(r'[^a-zA-Z0-9\s_]', '', clean_name)
    clean_name = clean_name.replace(' ', '_').lower()
    return clean_name

def map_type(col):
    t = col.get('type')
    size = ""
    dtype = ""
    fmt = ""
    if t == 'uuid':
        dtype = 'Varchar'
        size = '36 ký tự'
        fmt = 'Văn bản'
    elif t == 'text':
        dtype = 'Nvarchar'
        size = '255 ký tự'
        fmt = 'Văn bản'
        if col.get('multiline'):
            size = 'Văn bản dài'
            dtype = 'Text'
    elif t == 'integer':
        dtype = 'Int'
        size = ''
        fmt = 'Số nguyên dương' if ('min' in col and col['min'] >= 0) else 'Số nguyên'
    elif t == 'number':
        dtype = 'Decimal(18, 2)'
        size = ''
        fmt = 'Tiền tệ' if col.get('format') == 'currency_vnd' else 'Số thực'
    elif t == 'select':
        dtype = 'Nvarchar'
        size = '100 ký tự'
        fmt = 'Văn bản'
    elif t in ['date', 'datetime']:
        dtype = 'Datetime' if t == 'datetime' else 'Date'
        size = ''
        fmt = 'Thời gian'
    elif t == 'boolean':
        dtype = 'Bit'
        size = ''
        fmt = 'Đúng/Sai'
    elif t == 'phone':
        dtype = 'Varchar'
        size = '15 ký tự'
        fmt = 'Số điện thoại'
    elif t == 'email':
        dtype = 'Varchar'
        size = '100 ký tự'
        fmt = 'Email'
    elif t == 'url':
        dtype = 'Varchar'
        size = '255 ký tự'
        fmt = 'Đường dẫn (URL)'
    elif t == 'password':
        dtype = 'Varchar'
        size = '255 ký tự'
        fmt = 'Bảo mật'
    elif t == 'multi_select':
        dtype = 'Nvarchar'
        size = '255 ký tự'
        fmt = 'Văn bản'
    else:
        dtype = 'Nvarchar'
        size = '255 ký tự'
        fmt = 'Văn bản'

    if col.get('name').lower() == 'id' and t == 'integer':
         dtype = 'Int'
         fmt = 'Số nguyên dương'
         
    if dtype in ['Varchar', 'Nvarchar', 'Text']:
        fmt = 'Văn bản'
        
    return dtype, size, fmt

def create_word_doc():
    with open('d:/Dat/PBL3/Data/database_schema.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = docx.Document()
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.add_heading('Từ điển dữ liệu (Data Dictionary)', 0)

    # Colors
    HEADER_BG = 'ED7D31' # Orange
    EVEN_ROW_BG = 'FCE4D6' # Light peach
    ODD_ROW_BG = 'FFFFFF' # White

    for table_def in data['database']['tables']:
        doc.add_heading(f"Bảng: {table_def['name']} ({table_def['id']})", level=2)
        if table_def.get('description'):
            doc.add_paragraph(f"Mô tả: {table_def['description']}")

        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Define column widths
        widths = (Inches(1.2), Inches(1.0), Inches(0.8), Inches(1.2), Inches(1.0), Inches(1.8))
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Tên trường', 'Kiểu dữ liệu', 'Kích thước', 'Ràng buộc toàn vẹn', 'Khuôn dạng', 'Ghi chú']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].width = widths[i]
            set_cell_background(hdr_cells[i], HEADER_BG)
            # Center text and make it white & bold
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, col in enumerate(table_def['columns']):
            row = table.add_row()
            cells = row.cells
            
            name = to_snake_case(col['name'])
            if col.get('name') == 'ID' and col.get('type') == 'uuid':
                name = 'id'
                
            dtype, size, fmt = map_type(col)
            
            constraints = []
            if col.get('primary_key'): constraints.append('Primary key')
            if col.get('required') or col.get('primary_key'): constraints.append('Not null')
            if col.get('unique'): constraints.append('Unique')
            if col.get('foreign_key'): constraints.append('Foreign key')
            
            notes = []
            if 'options' in col: notes.append(f"Tùy chọn: {', '.join(col['options'])}")
            if 'default' in col: notes.append(f"Mặc định: {col['default']}")
            if col.get('references'):
                ref = col['references']
                notes.append(f"Tham chiếu: {ref['table']}.{ref['column']}")
            if col.get('auto_generate'): notes.append("Tự sinh")
            
            notes_str = ", ".join(notes) if notes else col.get('name')
            
            cells[0].text = name
            cells[0].paragraphs[0].runs[0].font.bold = True # Make field name bold
            cells[1].text = dtype
            cells[2].text = size
            cells[3].text = ", ".join(constraints)
            cells[4].text = fmt
            cells[5].text = notes_str

            # Apply background color
            bg_color = ODD_ROW_BG if row_idx % 2 == 0 else EVEN_ROW_BG
            for i, cell in enumerate(cells):
                cell.width = widths[i]
                set_cell_background(cell, bg_color)
                
                # Vertically center
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
                tcPr.append(vAlign)

                # Add some padding or just align text
                for paragraph in cell.paragraphs:
                    if i != 0 and i != 5: # center everything except name and notes usually, but let's just leave it left-aligned like image
                         pass

        doc.add_paragraph() # Add space between tables

    doc.save('d:/Dat/PBL3/Data/Data_Dictionary.docx')
    print("Done")

if __name__ == "__main__":
    create_word_doc()
